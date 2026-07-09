# -*- coding: utf-8 -*-
"""KAC 통역 평가 데이터 생성기.
   번역을 이 파일 한 곳에서 저작 → eval/dataset.json + eval/wordfiles/*.docx(언어별) 생성.
   실행: python3 eval/gen_dataset.py
   ⚠️ es/fr/pt/ar 은 초안 — 원어민 검수 후 확정할 것.
"""
import json, os

HERE = os.path.dirname(os.path.abspath(__file__))
FOREIGN = ['en', 'zh', 'ja', 'es', 'fr', 'pt', 'ar']
LANG_DISP = {'en': 'English', 'zh': '中文', 'ja': '日本語', 'es': 'Español', 'fr': 'Français', 'pt': 'Português', 'ar': 'العربية'}
LANG_KO = {'en': '영어', 'zh': '중국어', 'ja': '일본어', 'es': '스페인어', 'fr': '프랑스어', 'pt': '포르투갈어', 'ar': '아랍어'}
DIFF_KO = {'easy': '하', 'medium': '중', 'hard': '상'}
DIFF_ORDER = ['easy', 'medium', 'hard']

# 각 시나리오: q=질문(외국인 발화, 언어별) / q_ko=질문의 한국어 모범번역
#              a_ko=답변(한국인 발화) / a=답변의 언어별 모범번역
DATA = [
  # ===== 하 (easy) 10 =====
  {"id":"loc-restroom","cat":"위치","diff":"easy","traps":[],
   "q_ko":"화장실이 어디에 있나요?",
   "q":{"en":"Where is the restroom?","zh":"洗手间在哪里？","ja":"トイレはどこですか？","es":"¿Dónde está el baño?","fr":"Où sont les toilettes ?","pt":"Onde fica o banheiro?","ar":"أين دورة المياه؟"},
   "a_ko":"화장실은 3층 중앙 에스컬레이터 옆에 있습니다.",
   "a":{"en":"The restroom is next to the central escalator on the third floor.","zh":"洗手间在三楼中央自动扶梯旁边。","ja":"お手洗いは3階中央のエスカレーターの隣にあります。","es":"El baño está junto a la escalera mecánica central en el tercer piso.","fr":"Les toilettes se trouvent à côté de l'escalator central au troisième étage.","pt":"O banheiro fica ao lado da escada rolante central no terceiro andar.","ar":"دورة المياه بجانب السلّم الكهربائي المركزي في الطابق الثالث."}},

  {"id":"counter-korean-air","cat":"항공사","diff":"easy","traps":["proper_noun"],
   "q_ko":"대한항공 체크인 카운터는 어디에 있나요?",
   "q":{"en":"Where is the Korean Air check-in counter?","zh":"大韩航空的值机柜台在哪里？","ja":"大韓航空のチェックインカウンターはどこですか？","es":"¿Dónde está el mostrador de facturación de Korean Air?","fr":"Où se trouve le comptoir d'enregistrement de Korean Air ?","pt":"Onde fica o balcão de check-in da Korean Air?","ar":"أين مكتب تسجيل الوصول للخطوط الجوية الكورية (Korean Air)؟"},
   "a_ko":"대한항공 카운터는 3층 A구역에 있습니다.",
   "a":{"en":"The Korean Air counter is in Zone A on the third floor.","zh":"大韩航空柜台在三楼A区。","ja":"大韓航空のカウンターは3階のAエリアにあります。","es":"El mostrador de Korean Air está en la zona A del tercer piso.","fr":"Le comptoir de Korean Air se trouve dans la zone A au troisième étage.","pt":"O balcão da Korean Air fica na zona A, no terceiro andar.","ar":"مكتب الخطوط الجوية الكورية في المنطقة A بالطابق الثالث."}},

  {"id":"loc-atm","cat":"편의","diff":"easy","traps":[],
   "q_ko":"현금인출기(ATM)는 어디에 있나요?",
   "q":{"en":"Where is the ATM?","zh":"自动取款机（ATM）在哪里？","ja":"ATM（現金自動預払機）はどこですか？","es":"¿Dónde está el cajero automático?","fr":"Où se trouve le distributeur automatique de billets ?","pt":"Onde fica o caixa eletrônico?","ar":"أين جهاز الصراف الآلي؟"},
   "a_ko":"ATM은 2층 편의점 옆에 있습니다.",
   "a":{"en":"The ATM is next to the convenience store on the second floor.","zh":"自动取款机在二楼便利店旁边。","ja":"ATMは2階のコンビニの隣にあります。","es":"El cajero automático está junto a la tienda de conveniencia en el segundo piso.","fr":"Le distributeur se trouve à côté de la supérette au deuxième étage.","pt":"O caixa eletrônico fica ao lado da loja de conveniência no segundo andar.","ar":"جهاز الصراف الآلي بجانب المتجر الصغير في الطابق الثاني."}},

  {"id":"loc-convenience","cat":"편의","diff":"easy","traps":[],
   "q_ko":"편의점은 어디에 있나요?",
   "q":{"en":"Where is the convenience store?","zh":"便利店在哪里？","ja":"コンビニはどこですか？","es":"¿Dónde está la tienda de conveniencia?","fr":"Où se trouve la supérette ?","pt":"Onde fica a loja de conveniência?","ar":"أين المتجر الصغير؟"},
   "a_ko":"편의점은 1층 입국장 옆에 있습니다.",
   "a":{"en":"The convenience store is next to the arrivals hall on the first floor.","zh":"便利店在一楼入境大厅旁边。","ja":"コンビニは1階の到着ロビーの隣にあります。","es":"La tienda de conveniencia está junto a la sala de llegadas en el primer piso.","fr":"La supérette se trouve à côté du hall des arrivées au premier étage.","pt":"A loja de conveniência fica ao lado do saguão de desembarque no primeiro andar.","ar":"المتجر الصغير بجانب صالة الوصول في الطابق الأول."}},

  {"id":"counter-asiana","cat":"항공사","diff":"easy","traps":["proper_noun"],
   "q_ko":"아시아나항공 카운터는 몇 층인가요?",
   "q":{"en":"What floor is the Asiana Airlines counter on?","zh":"韩亚航空的柜台在几楼？","ja":"アシアナ航空のカウンターは何階ですか？","es":"¿En qué piso está el mostrador de Asiana Airlines?","fr":"À quel étage se trouve le comptoir d'Asiana Airlines ?","pt":"Em que andar fica o balcão da Asiana Airlines?","ar":"في أي طابق يقع مكتب الخطوط الجوية الآسيوية (Asiana Airlines)؟"},
   "a_ko":"아시아나항공 카운터는 3층 L구역에 있습니다.",
   "a":{"en":"The Asiana Airlines counter is in Zone L on the third floor.","zh":"韩亚航空柜台在三楼L区。","ja":"アシアナ航空のカウンターは3階のLエリアにあります。","es":"El mostrador de Asiana Airlines está en la zona L del tercer piso.","fr":"Le comptoir d'Asiana Airlines se trouve dans la zone L au troisième étage.","pt":"O balcão da Asiana Airlines fica na zona L, no terceiro andar.","ar":"مكتب الخطوط الجوية الآسيوية في المنطقة L بالطابق الثالث."}},

  {"id":"loc-smoking","cat":"편의","diff":"easy","traps":[],
   "q_ko":"흡연실은 어디에 있나요?",
   "q":{"en":"Where is the smoking room?","zh":"吸烟室在哪里？","ja":"喫煙室はどこですか？","es":"¿Dónde está la sala de fumadores?","fr":"Où se trouve le fumoir ?","pt":"Onde fica a sala de fumantes?","ar":"أين غرفة التدخين؟"},
   "a_ko":"흡연실은 4층 야외 테라스에 있습니다.",
   "a":{"en":"The smoking room is on the outdoor terrace on the fourth floor.","zh":"吸烟室在四楼室外露台。","ja":"喫煙室は4階の屋外テラスにあります。","es":"La sala de fumadores está en la terraza exterior del cuarto piso.","fr":"Le fumoir se trouve sur la terrasse extérieure au quatrième étage.","pt":"A sala de fumantes fica no terraço externo do quarto andar.","ar":"غرفة التدخين في الشرفة الخارجية بالطابق الرابع."}},

  {"id":"loc-charging","cat":"편의","diff":"easy","traps":[],
   "q_ko":"휴대폰을 충전할 수 있는 곳이 어디인가요?",
   "q":{"en":"Where can I charge my phone?","zh":"哪里可以给手机充电？","ja":"携帯電話を充電できる場所はどこですか？","es":"¿Dónde puedo cargar mi teléfono?","fr":"Où puis-je recharger mon téléphone ?","pt":"Onde posso carregar meu celular?","ar":"أين يمكنني شحن هاتفي؟"},
   "a_ko":"충전 구역은 각 게이트 앞 대기 의자에 있습니다.",
   "a":{"en":"Charging stations are at the waiting seats in front of each gate.","zh":"充电区在每个登机口前的候机座位处。","ja":"充電コーナーは各搭乗口前の待合席にあります。","es":"Las estaciones de carga están en los asientos de espera frente a cada puerta.","fr":"Les bornes de recharge se trouvent aux sièges d'attente devant chaque porte.","pt":"Os pontos de recarga ficam nos assentos de espera em frente a cada portão.","ar":"نقاط الشحن موجودة عند مقاعد الانتظار أمام كل بوابة."}},

  {"id":"loc-taxi","cat":"교통","diff":"easy","traps":[],
   "q_ko":"택시는 어디에서 타나요?",
   "q":{"en":"Where do I catch a taxi?","zh":"在哪里乘坐出租车？","ja":"タクシーはどこで乗れますか？","es":"¿Dónde puedo tomar un taxi?","fr":"Où puis-je prendre un taxi ?","pt":"Onde posso pegar um táxi?","ar":"من أين آخذ سيارة أجرة؟"},
   "a_ko":"택시 승강장은 1층 4번 출구 앞에 있습니다.",
   "a":{"en":"The taxi stand is in front of Exit 4 on the first floor.","zh":"出租车乘车点在一楼4号出口前。","ja":"タクシー乗り場は1階の4番出口の前にあります。","es":"La parada de taxis está frente a la salida 4 en el primer piso.","fr":"La station de taxis se trouve devant la sortie 4 au premier étage.","pt":"O ponto de táxi fica em frente à saída 4 no primeiro andar.","ar":"موقف سيارات الأجرة أمام المخرج رقم 4 في الطابق الأول."}},

  {"id":"counter-jeju-air","cat":"항공사","diff":"easy","traps":["proper_noun"],
   "q_ko":"제주항공 카운터는 어디에 있나요?",
   "q":{"en":"Where is the Jeju Air counter?","zh":"济州航空的柜台在哪里？","ja":"チェジュ航空のカウンターはどこですか？","es":"¿Dónde está el mostrador de Jeju Air?","fr":"Où se trouve le comptoir de Jeju Air ?","pt":"Onde fica o balcão da Jeju Air?","ar":"أين مكتب خطوط جيجو الجوية (Jeju Air)؟"},
   "a_ko":"제주항공 카운터는 3층 D구역에 있습니다.",
   "a":{"en":"The Jeju Air counter is in Zone D on the third floor.","zh":"济州航空柜台在三楼D区。","ja":"チェジュ航空のカウンターは3階のDエリアにあります。","es":"El mostrador de Jeju Air está en la zona D del tercer piso.","fr":"Le comptoir de Jeju Air se trouve dans la zone D au troisième étage.","pt":"O balcão da Jeju Air fica na zona D, no terceiro andar.","ar":"مكتب خطوط جيجو الجوية في المنطقة D بالطابق الثالث."}},

  {"id":"loc-exit","cat":"위치","diff":"easy","traps":[],
   "q_ko":"가장 가까운 출구는 어디인가요?",
   "q":{"en":"Where is the nearest exit?","zh":"最近的出口在哪里？","ja":"一番近い出口はどこですか？","es":"¿Dónde está la salida más cercana?","fr":"Où est la sortie la plus proche ?","pt":"Onde fica a saída mais próxima?","ar":"أين أقرب مخرج؟"},
   "a_ko":"가장 가까운 출구는 왼쪽으로 50미터 앞 2번 출구입니다.",
   "a":{"en":"The nearest exit is Exit 2, 50 meters ahead on your left.","zh":"最近的出口是左前方50米处的2号出口。","ja":"一番近い出口は、左手前方50メートルの2番出口です。","es":"La salida más cercana es la salida 2, a 50 metros a su izquierda.","fr":"La sortie la plus proche est la sortie 2, à 50 mètres sur votre gauche.","pt":"A saída mais próxima é a saída 2, a 50 metros à sua esquerda.","ar":"أقرب مخرج هو المخرج رقم 2، على بعد 50 مترًا إلى يسارك."}},

  # ===== 중 (medium) 10 =====
  {"id":"gate-find","cat":"위치","diff":"medium","traps":["gate_code","number"],
   "q_ko":"27번 게이트에 어떻게 가나요?",
   "q":{"en":"How do I get to Gate 27?","zh":"27号登机口怎么走？","ja":"27番搭乗口へはどう行けばいいですか？","es":"¿Cómo llego a la puerta 27?","fr":"Comment aller à la porte 27 ?","pt":"Como chego ao portão 27?","ar":"كيف أصل إلى البوابة 27؟"},
   "a_ko":"27번 게이트는 이 복도를 따라가면 오른쪽 끝에 있습니다.",
   "a":{"en":"Gate 27 is at the far right end if you follow this hallway.","zh":"沿着这条走廊走，27号登机口在右侧尽头。","ja":"この廊下に沿って進むと、27番搭乗口は右側の突き当たりにあります。","es":"La puerta 27 está al fondo a la derecha si sigue este pasillo.","fr":"La porte 27 se trouve tout au bout à droite en suivant ce couloir.","pt":"O portão 27 fica no fim à direita se você seguir por este corredor.","ar":"البوابة 27 في نهاية هذا الممر على الجهة اليمنى."}},

  {"id":"transfer-counter","cat":"절차","diff":"medium","traps":["procedure","proper_noun"],
   "q_ko":"환승 카운터는 어디에 있나요?",
   "q":{"en":"Where is the transfer counter?","zh":"转机柜台在哪里？","ja":"乗り継ぎカウンターはどこですか？","es":"¿Dónde está el mostrador de tránsito?","fr":"Où se trouve le comptoir de correspondance ?","pt":"Onde fica o balcão de conexão?","ar":"أين مكتب الترانزيت؟"},
   "a_ko":"환승 카운터는 한 층 위 면세구역 입구에 있습니다.",
   "a":{"en":"The transfer counter is one floor up, at the entrance to the duty-free area.","zh":"转机柜台在上一层，免税区入口处。","ja":"乗り継ぎカウンターは1つ上の階、免税エリアの入口にあります。","es":"El mostrador de tránsito está un piso más arriba, en la entrada de la zona libre de impuestos.","fr":"Le comptoir de correspondance est à l'étage au-dessus, à l'entrée de la zone hors taxes.","pt":"O balcão de conexão fica um andar acima, na entrada da área duty-free.","ar":"مكتب الترانزيت في الطابق الذي يعلوه، عند مدخل منطقة السوق الحرة."}},

  {"id":"baggage-wrap","cat":"수하물","diff":"medium","traps":[],
   "q_ko":"수하물 포장 서비스는 어디에서 하나요?",
   "q":{"en":"Where can I get my baggage wrapped?","zh":"在哪里可以打包行李？","ja":"手荷物の梱包サービスはどこでできますか？","es":"¿Dónde puedo envolver mi equipaje?","fr":"Où puis-je faire emballer mes bagages ?","pt":"Onde posso embalar minha bagagem?","ar":"أين يمكنني تغليف أمتعتي؟"},
   "a_ko":"수하물 포장 서비스는 3층 체크인 카운터 양 끝에 있습니다.",
   "a":{"en":"The baggage wrapping service is at both ends of the check-in counters on the third floor.","zh":"行李打包服务在三楼值机柜台的两端。","ja":"手荷物梱包サービスは3階のチェックインカウンターの両端にあります。","es":"El servicio de envoltura de equipaje está en ambos extremos de los mostradores de facturación del tercer piso.","fr":"Le service d'emballage des bagages se trouve aux deux extrémités des comptoirs d'enregistrement au troisième étage.","pt":"O serviço de embalagem de bagagem fica nas duas extremidades dos balcões de check-in no terceiro andar.","ar":"خدمة تغليف الأمتعة موجودة على طرفي مكاتب تسجيل الوصول في الطابق الثالث."}},

  {"id":"checkin-time-jinair","cat":"항공사","diff":"medium","traps":["proper_noun","time"],
   "q_ko":"진에어 체크인은 언제 시작하나요?",
   "q":{"en":"When does Jin Air check-in open?","zh":"真航空的值机什么时候开始？","ja":"ジンエアーのチェックインはいつ始まりますか？","es":"¿Cuándo abre la facturación de Jin Air?","fr":"Quand ouvre l'enregistrement de Jin Air ?","pt":"Quando abre o check-in da Jin Air?","ar":"متى يبدأ تسجيل الوصول لطيران جين (Jin Air)؟"},
   "a_ko":"진에어 체크인은 보통 출발 2시간 30분 전에 시작합니다.",
   "a":{"en":"Jin Air check-in usually opens 2 hours and 30 minutes before departure.","zh":"真航空的值机通常在起飞前2小时30分钟开始。","ja":"ジンエアーのチェックインは通常、出発の2時間30分前に始まります。","es":"La facturación de Jin Air suele abrir 2 horas y 30 minutos antes de la salida.","fr":"L'enregistrement de Jin Air ouvre généralement 2 heures 30 avant le départ.","pt":"O check-in da Jin Air geralmente abre 2 horas e 30 minutos antes da partida.","ar":"عادةً يبدأ تسجيل الوصول لطيران جين قبل المغادرة بساعتين ونصف."}},

  {"id":"bus-terminal","cat":"교통","diff":"medium","traps":["number"],
   "q_ko":"시내로 가는 공항버스는 어디에서 타나요?",
   "q":{"en":"Where do I catch the airport bus to the city?","zh":"去市区的机场大巴在哪里乘坐？","ja":"市内へ行く空港バスはどこで乗れますか？","es":"¿Dónde tomo el autobús del aeropuerto hacia la ciudad?","fr":"Où puis-je prendre le bus de l'aéroport pour aller en ville ?","pt":"Onde pego o ônibus do aeroporto para a cidade?","ar":"من أين آخذ حافلة المطار المتجهة إلى المدينة؟"},
   "a_ko":"공항버스 정류장은 1층 밖으로 나가서 4번과 6번 승강장에 있습니다.",
   "a":{"en":"The airport bus stops are outside on the first floor, at platforms 4 and 6.","zh":"机场大巴车站在一楼外面的4号和6号站台。","ja":"空港バス乗り場は1階の外に出て、4番と6番のりばにあります。","es":"Las paradas del autobús del aeropuerto están afuera en el primer piso, en los andenes 4 y 6.","fr":"Les arrêts du bus de l'aéroport se trouvent à l'extérieur au premier étage, aux quais 4 et 6.","pt":"Os pontos do ônibus do aeroporto ficam do lado de fora no primeiro andar, nas plataformas 4 e 6.","ar":"مواقف حافلات المطار في الخارج بالطابق الأول، عند الرصيفين 4 و6."}},

  {"id":"lost-item","cat":"편의","diff":"medium","traps":[],
   "q_ko":"분실물 센터는 어디에 있나요?",
   "q":{"en":"Where is the lost and found?","zh":"失物招领处在哪里？","ja":"忘れ物センターはどこですか？","es":"¿Dónde está la oficina de objetos perdidos?","fr":"Où se trouve le bureau des objets trouvés ?","pt":"Onde fica o achados e perdidos?","ar":"أين مكتب المفقودات؟"},
   "a_ko":"분실물 센터는 2층 안내데스크 뒤편에 있습니다.",
   "a":{"en":"The lost and found is behind the information desk on the second floor.","zh":"失物招领处在二楼问询处后面。","ja":"忘れ物センターは2階の案内デスクの裏側にあります。","es":"La oficina de objetos perdidos está detrás del mostrador de información en el segundo piso.","fr":"Le bureau des objets trouvés se trouve derrière le comptoir d'information au deuxième étage.","pt":"O achados e perdidos fica atrás do balcão de informações no segundo andar.","ar":"مكتب المفقودات خلف مكتب الاستعلامات في الطابق الثاني."}},

  {"id":"prayer-room","cat":"편의","diff":"medium","traps":[],
   "q_ko":"기도실은 어디에 있나요?",
   "q":{"en":"Where is the prayer room?","zh":"祈祷室在哪里？","ja":"祈祷室はどこですか？","es":"¿Dónde está la sala de oración?","fr":"Où se trouve la salle de prière ?","pt":"Onde fica a sala de oração?","ar":"أين غرفة الصلاة؟"},
   "a_ko":"기도실은 4층 문화센터 옆에 있습니다.",
   "a":{"en":"The prayer room is next to the cultural center on the fourth floor.","zh":"祈祷室在四楼文化中心旁边。","ja":"祈祷室は4階の文化センターの隣にあります。","es":"La sala de oración está junto al centro cultural en el cuarto piso.","fr":"La salle de prière se trouve à côté du centre culturel au quatrième étage.","pt":"A sala de oração fica ao lado do centro cultural no quarto andar.","ar":"غرفة الصلاة بجانب المركز الثقافي في الطابق الرابع."}},

  {"id":"duty-free-pickup","cat":"절차","diff":"medium","traps":["gate_code"],
   "q_ko":"면세품 인도장은 어디에 있나요?",
   "q":{"en":"Where is the duty-free pickup counter?","zh":"免税品提货处在哪里？","ja":"免税品の受け取り場所はどこですか？","es":"¿Dónde está el punto de recogida de artículos libres de impuestos?","fr":"Où se trouve le point de retrait des achats hors taxes ?","pt":"Onde fica o balcão de retirada de produtos duty-free?","ar":"أين نقطة استلام مشتريات السوق الحرة؟"},
   "a_ko":"면세품 인도장은 출국 심사를 지나 28번 게이트 근처에 있습니다.",
   "a":{"en":"The duty-free pickup counter is near Gate 28, after passing through immigration.","zh":"免税品提货处在过了出境检查后的28号登机口附近。","ja":"免税品の受け取り場所は、出国審査を過ぎた28番搭乗口の近くにあります。","es":"El punto de recogida de duty-free está cerca de la puerta 28, después de pasar el control de inmigración.","fr":"Le point de retrait hors taxes se trouve près de la porte 28, après le contrôle de l'immigration.","pt":"O balcão de retirada duty-free fica perto do portão 28, depois de passar pela imigração.","ar":"نقطة استلام مشتريات السوق الحرة قرب البوابة 28، بعد اجتياز مراقبة الجوازات."}},

  {"id":"counter-tway","cat":"수하물","diff":"medium","traps":["proper_noun"],
   "q_ko":"티웨이항공 수하물 규정을 어디에서 물어보나요?",
   "q":{"en":"Where can I ask about T'way Air baggage rules?","zh":"在哪里可以咨询德威航空的行李规定？","ja":"ティーウェイ航空の手荷物規定はどこで聞けますか？","es":"¿Dónde puedo preguntar sobre las normas de equipaje de T'way Air?","fr":"Où puis-je me renseigner sur les règles de bagages de T'way Air ?","pt":"Onde posso perguntar sobre as regras de bagagem da T'way Air?","ar":"أين يمكنني الاستفسار عن قواعد الأمتعة لطيران تيوي (T'way Air)؟"},
   "a_ko":"티웨이항공 카운터는 3층 H구역에 있으니 그곳에서 안내받으실 수 있습니다.",
   "a":{"en":"The T'way Air counter is in Zone H on the third floor, where you can get assistance.","zh":"德威航空柜台在三楼H区，您可以在那里咨询。","ja":"ティーウェイ航空のカウンターは3階のHエリアにありますので、そちらでご案内を受けられます。","es":"El mostrador de T'way Air está en la zona H del tercer piso, donde pueden atenderle.","fr":"Le comptoir de T'way Air se trouve dans la zone H au troisième étage, où l'on pourra vous renseigner.","pt":"O balcão da T'way Air fica na zona H, no terceiro andar, onde você pode ser atendido.","ar":"مكتب طيران تيوي في المنطقة H بالطابق الثالث، حيث يمكنك الحصول على المساعدة."}},

  {"id":"wheelchair","cat":"편의","diff":"medium","traps":[],
   "q_ko":"휠체어 도움을 어디에서 요청하나요?",
   "q":{"en":"Where can I request wheelchair assistance?","zh":"在哪里可以申请轮椅协助？","ja":"車椅子のサポートはどこで頼めますか？","es":"¿Dónde puedo solicitar asistencia en silla de ruedas?","fr":"Où puis-je demander une assistance en fauteuil roulant ?","pt":"Onde posso solicitar assistência de cadeira de rodas?","ar":"أين يمكنني طلب مساعدة الكرسي المتحرك؟"},
   "a_ko":"휠체어 지원은 1층 안내데스크에서 요청하실 수 있습니다.",
   "a":{"en":"You can request wheelchair assistance at the information desk on the first floor.","zh":"您可以在一楼问询处申请轮椅协助。","ja":"車椅子のサポートは1階の案内デスクでお申し込みいただけます。","es":"Puede solicitar asistencia en silla de ruedas en el mostrador de información del primer piso.","fr":"Vous pouvez demander une assistance en fauteuil roulant au comptoir d'information au premier étage.","pt":"Você pode solicitar assistência de cadeira de rodas no balcão de informações no primeiro andar.","ar":"يمكنك طلب مساعدة الكرسي المتحرك من مكتب الاستعلامات في الطابق الأول."}},

  # ===== 상 (hard) 10 =====
  {"id":"baggage-fee","cat":"수하물","diff":"hard","traps":["number","currency"],
   "q_ko":"초과 수하물 요금은 얼마인가요?",
   "q":{"en":"How much is the excess baggage fee?","zh":"超重行李费是多少？","ja":"超過手荷物の料金はいくらですか？","es":"¿Cuánto cuesta el exceso de equipaje?","fr":"Combien coûtent les frais d'excédent de bagages ?","pt":"Quanto custa a taxa de excesso de bagagem?","ar":"كم رسوم الوزن الزائد للأمتعة؟"},
   "a_ko":"초과 수하물 요금은 1킬로그램당 만 오천 원입니다.",
   "a":{"en":"The excess baggage fee is 15,000 won per kilogram.","zh":"超重行李费为每公斤一万五千韩元。","ja":"超過手荷物の料金は1キログラムあたり15,000ウォンです。","es":"La tarifa por exceso de equipaje es de 15.000 wones por kilogramo.","fr":"Les frais d'excédent de bagages sont de 15 000 wons par kilogramme.","pt":"A taxa de excesso de bagagem é de 15.000 wones por quilograma.","ar":"رسوم الوزن الزائد للأمتعة هي 15,000 وون لكل كيلوغرام."}},

  {"id":"flight-delay","cat":"시간·편명","diff":"hard","traps":["flight_code","time","number"],
   "q_ko":"KE123편이 지연되나요?",
   "q":{"en":"Is flight KE123 delayed?","zh":"KE123航班延误了吗？","ja":"KE123便は遅れていますか？","es":"¿El vuelo KE123 está retrasado?","fr":"Le vol KE123 est-il retardé ?","pt":"O voo KE123 está atrasado?","ar":"هل الرحلة KE123 متأخرة؟"},
   "a_ko":"KE123편은 30분 지연되어 오후 2시 40분에 출발합니다.",
   "a":{"en":"Flight KE123 is delayed by 30 minutes and will depart at 2:40 PM.","zh":"KE123航班延误30分钟，将于下午2点40分起飞。","ja":"KE123便は30分遅れて、午後2時40分に出発します。","es":"El vuelo KE123 está retrasado 30 minutos y saldrá a las 14:40.","fr":"Le vol KE123 est retardé de 30 minutes et partira à 14h40.","pt":"O voo KE123 está atrasado 30 minutos e partirá às 14h40.","ar":"الرحلة KE123 متأخرة 30 دقيقة وستغادر في الساعة 2:40 مساءً."}},

  {"id":"flight-status-koreanair","cat":"항공사","diff":"hard","traps":["flight_code","proper_noun"],
   "q_ko":"대한항공 KE081편의 탑승 시간을 알 수 있나요?",
   "q":{"en":"Can you tell me the boarding time for Korean Air flight KE081?","zh":"能告诉我大韩航空KE081航班的登机时间吗？","ja":"大韓航空KE081便の搭乗時刻を教えてもらえますか？","es":"¿Me puede decir la hora de embarque del vuelo KE081 de Korean Air?","fr":"Pouvez-vous me dire l'heure d'embarquement du vol KE081 de Korean Air ?","pt":"Pode me dizer o horário de embarque do voo KE081 da Korean Air?","ar":"هل يمكنك إخباري بوقت الصعود للرحلة KE081 للخطوط الجوية الكورية؟"},
   "a_ko":"정확한 탑승 시간은 대한항공 카운터에서 확인해 주세요. 3층 A구역에 있습니다.",
   "a":{"en":"Please check the exact boarding time at the Korean Air counter. It is in Zone A on the third floor.","zh":"准确的登机时间请在大韩航空柜台确认。柜台在三楼A区。","ja":"正確な搭乗時刻は大韓航空のカウンターでご確認ください。3階のAエリアにあります。","es":"Consulte la hora exacta de embarque en el mostrador de Korean Air. Está en la zona A del tercer piso.","fr":"Veuillez vérifier l'heure exacte d'embarquement au comptoir de Korean Air. Il se trouve dans la zone A au troisième étage.","pt":"Por favor, confirme o horário exato de embarque no balcão da Korean Air. Fica na zona A, no terceiro andar.","ar":"يرجى التحقق من وقت الصعود الدقيق لدى مكتب الخطوط الجوية الكورية. يقع في المنطقة A بالطابق الثالث."}},

  {"id":"connect-time-asiana","cat":"항공사","diff":"hard","traps":["flight_code","proper_noun","procedure"],
   "q_ko":"아시아나 OZ102편으로 갈아타는데 시간이 촉박해요. 어떻게 하나요?",
   "q":{"en":"I have a tight connection to Asiana flight OZ102. What should I do?","zh":"我要转乘韩亚航空OZ102航班，时间很紧，该怎么办？","ja":"アシアナ航空OZ102便に乗り継ぐのですが、時間が迫っています。どうすればいいですか？","es":"Tengo una conexión muy justa con el vuelo OZ102 de Asiana. ¿Qué debo hacer?","fr":"J'ai une correspondance très serrée pour le vol OZ102 d'Asiana. Que dois-je faire ?","pt":"Tenho uma conexão apertada para o voo OZ102 da Asiana. O que devo fazer?","ar":"لدي وقت ضيق للتوصيلة إلى رحلة الخطوط الآسيوية OZ102. ماذا أفعل؟"},
   "a_ko":"자세한 환승 안내는 아시아나항공 카운터에서 도와드립니다. 환승 구역 안쪽에 있습니다.",
   "a":{"en":"For detailed transfer assistance, please go to the Asiana Airlines counter. It is inside the transfer area.","zh":"详细的转机指引请到韩亚航空柜台咨询。柜台在转机区里面。","ja":"詳しい乗り継ぎのご案内はアシアナ航空のカウンターでお手伝いします。乗り継ぎエリアの内側にあります。","es":"Para asistencia detallada de conexión, diríjase al mostrador de Asiana Airlines. Está dentro de la zona de tránsito.","fr":"Pour une assistance détaillée de correspondance, rendez-vous au comptoir d'Asiana Airlines. Il se trouve à l'intérieur de la zone de correspondance.","pt":"Para assistência detalhada de conexão, dirija-se ao balcão da Asiana Airlines. Fica dentro da área de conexão.","ar":"للحصول على مساعدة مفصّلة بشأن التوصيلة، توجّه إلى مكتب الخطوط الجوية الآسيوية. يقع داخل منطقة الترانزيت."}},

  {"id":"baggage-lost-jeju","cat":"수하물","diff":"hard","traps":["proper_noun","procedure"],
   "q_ko":"제주항공편으로 왔는데 수하물이 나오지 않았어요.",
   "q":{"en":"I arrived on a Jeju Air flight, but my baggage didn't come out.","zh":"我乘坐济州航空的航班到达，但我的行李没有出来。","ja":"チェジュ航空の便で到着しましたが、手荷物が出てきませんでした。","es":"Llegué en un vuelo de Jeju Air, pero mi equipaje no salió.","fr":"Je suis arrivé sur un vol de Jeju Air, mais mes bagages ne sont pas sortis.","pt":"Cheguei em um voo da Jeju Air, mas minha bagagem não saiu.","ar":"وصلت على متن رحلة لخطوط جيجو الجوية، لكن أمتعتي لم تخرج."},
   "a_ko":"수하물 분실은 제주항공 수하물 카운터에서 접수해 주세요. 수하물 수취대 옆에 있습니다.",
   "a":{"en":"Please report the lost baggage at the Jeju Air baggage counter. It is next to the baggage claim.","zh":"行李丢失请在济州航空行李柜台登记。柜台在行李提取处旁边。","ja":"手荷物の紛失はチェジュ航空の手荷物カウンターでお手続きください。手荷物受取所の隣にあります。","es":"Por favor, reporte el equipaje perdido en el mostrador de equipaje de Jeju Air. Está junto a la cinta de equipaje.","fr":"Veuillez signaler les bagages perdus au comptoir des bagages de Jeju Air. Il se trouve à côté du tapis à bagages.","pt":"Por favor, registre a bagagem perdida no balcão de bagagem da Jeju Air. Fica ao lado da esteira de bagagem.","ar":"يرجى الإبلاغ عن الأمتعة المفقودة لدى مكتب أمتعة خطوط جيجو الجوية. يقع بجانب سير استلام الأمتعة."}},

  {"id":"refund-cancel-tway","cat":"항공사","diff":"hard","traps":["proper_noun","procedure"],
   "q_ko":"티웨이항공 항공편이 취소됐는데 환불은 어떻게 받나요?",
   "q":{"en":"My T'way Air flight was canceled. How do I get a refund?","zh":"我的德威航空航班取消了，怎么退款？","ja":"ティーウェイ航空の便が欠航になったのですが、払い戻しはどうすればいいですか？","es":"Mi vuelo de T'way Air fue cancelado. ¿Cómo obtengo un reembolso?","fr":"Mon vol T'way Air a été annulé. Comment obtenir un remboursement ?","pt":"Meu voo da T'way Air foi cancelado. Como consigo o reembolso?","ar":"أُلغيت رحلتي على طيران تيوي. كيف أحصل على استرداد المبلغ؟"},
   "a_ko":"환불과 재예약은 티웨이항공 카운터에서 처리해 드립니다. 3층 H구역입니다.",
   "a":{"en":"Refunds and rebooking are handled at the T'way Air counter. It is in Zone H on the third floor.","zh":"退款和改签由德威航空柜台办理。柜台在三楼H区。","ja":"払い戻しと予約の取り直しはティーウェイ航空のカウンターで承ります。3階のHエリアです。","es":"Los reembolsos y las reprogramaciones se gestionan en el mostrador de T'way Air. Está en la zona H del tercer piso.","fr":"Les remboursements et les nouvelles réservations se font au comptoir de T'way Air. Il se trouve dans la zone H au troisième étage.","pt":"Reembolsos e remarcações são feitos no balcão da T'way Air. Fica na zona H, no terceiro andar.","ar":"تتم عمليات الاسترداد وإعادة الحجز لدى مكتب طيران تيوي. يقع في المنطقة H بالطابق الثالث."}},

  {"id":"oversize-sports","cat":"수하물","diff":"hard","traps":["number"],
   "q_ko":"골프백은 추가 요금이 있나요?",
   "q":{"en":"Is there an extra charge for a golf bag?","zh":"高尔夫球包需要额外收费吗？","ja":"ゴルフバッグは追加料金がかかりますか？","es":"¿Hay un cargo adicional por una bolsa de golf?","fr":"Y a-t-il un supplément pour un sac de golf ?","pt":"Há uma taxa extra para uma bolsa de golfe?","ar":"هل هناك رسوم إضافية لحقيبة الغولف؟"},
   "a_ko":"골프백 같은 스포츠 장비는 별도 요금이 있으며, 크기와 무게에 따라 다릅니다.",
   "a":{"en":"Sports equipment like golf bags has a separate charge, which varies by size and weight.","zh":"高尔夫球包等运动装备需另外收费，具体金额根据尺寸和重量而定。","ja":"ゴルフバッグなどのスポーツ用品は別料金がかかり、サイズと重さによって異なります。","es":"El equipo deportivo, como las bolsas de golf, tiene un cargo aparte que varía según el tamaño y el peso.","fr":"Les équipements sportifs comme les sacs de golf font l'objet d'un supplément qui varie selon la taille et le poids.","pt":"Equipamentos esportivos, como bolsas de golfe, têm uma taxa à parte, que varia conforme o tamanho e o peso.","ar":"المعدات الرياضية مثل حقائب الغولف لها رسوم منفصلة تختلف حسب الحجم والوزن."}},

  {"id":"medication-liquid","cat":"절차","diff":"hard","traps":["number","procedure"],
   "q_ko":"액체로 된 의약품을 기내에 가지고 탈 수 있나요?",
   "q":{"en":"Can I bring liquid medication in my carry-on?","zh":"我可以把液体药品带上飞机吗？","ja":"液体の医薬品を機内に持ち込めますか？","es":"¿Puedo llevar medicamentos líquidos en el equipaje de mano?","fr":"Puis-je emporter des médicaments liquides dans mon bagage à main ?","pt":"Posso levar medicamentos líquidos na bagagem de mão?","ar":"هل يمكنني إحضار أدوية سائلة في أمتعة اليد؟"},
   "a_ko":"의약품은 100밀리리터를 넘어도 반입할 수 있지만, 처방전이나 증빙을 보안검색대에서 보여 주셔야 합니다.",
   "a":{"en":"Medication is allowed even over 100 milliliters, but you must show a prescription or proof at the security checkpoint.","zh":"药品即使超过100毫升也可以带入，但您需要在安检处出示处方或证明。","ja":"医薬品は100ミリリットルを超えても持ち込めますが、処方箋や証明書を保安検査場で提示する必要があります。","es":"Los medicamentos están permitidos incluso si superan los 100 mililitros, pero debe mostrar una receta o comprobante en el control de seguridad.","fr":"Les médicaments sont autorisés même au-delà de 100 millilitres, mais vous devez présenter une ordonnance ou un justificatif au contrôle de sécurité.","pt":"Medicamentos são permitidos mesmo acima de 100 mililitros, mas você deve apresentar uma receita ou comprovante no controle de segurança.","ar":"يُسمح بالأدوية حتى لو تجاوزت 100 مليلتر، لكن يجب إظهار وصفة طبية أو إثبات عند نقطة التفتيش الأمني."}},

  {"id":"gate-change-airbusan","cat":"항공사","diff":"hard","traps":["flight_code","proper_noun"],
   "q_ko":"에어부산 BX 항공편 게이트가 변경됐나요?",
   "q":{"en":"Has the gate changed for my Air Busan (BX) flight?","zh":"我的釜山航空（BX）航班登机口有变化吗？","ja":"エアプサン（BX）便の搭乗口は変更されましたか？","es":"¿Ha cambiado la puerta de mi vuelo de Air Busan (BX)?","fr":"La porte de mon vol Air Busan (BX) a-t-elle changé ?","pt":"O portão do meu voo da Air Busan (BX) mudou?","ar":"هل تغيّرت بوابة رحلتي على طيران بوسان (BX)؟"},
   "a_ko":"게이트 변경은 에어부산 카운터나 전광판에서 확인해 주세요. 카운터는 3층 G구역에 있습니다.",
   "a":{"en":"Please check for gate changes at the Air Busan counter or on the display board. The counter is in Zone G on the third floor.","zh":"登机口变更请在釜山航空柜台或航班信息屏确认。柜台在三楼G区。","ja":"搭乗口の変更はエアプサンのカウンターまたは電光掲示板でご確認ください。カウンターは3階のGエリアにあります。","es":"Por favor, verifique los cambios de puerta en el mostrador de Air Busan o en el panel de información. El mostrador está en la zona G del tercer piso.","fr":"Veuillez vérifier les changements de porte au comptoir d'Air Busan ou sur le panneau d'affichage. Le comptoir se trouve dans la zone G au troisième étage.","pt":"Por favor, verifique as mudanças de portão no balcão da Air Busan ou no painel de informações. O balcão fica na zona G, no terceiro andar.","ar":"يرجى التحقق من تغييرات البوابة لدى مكتب طيران بوسان أو على شاشة العرض. يقع المكتب في المنطقة G بالطابق الثالث."}},

  {"id":"customs-limit","cat":"절차","diff":"hard","traps":["currency","number"],
   "q_ko":"면세 한도는 얼마까지인가요?",
   "q":{"en":"What is the duty-free allowance limit?","zh":"免税额度是多少？","ja":"免税の範囲はいくらまでですか？","es":"¿Cuál es el límite de la franquicia libre de impuestos?","fr":"Quelle est la limite de la franchise hors taxes ?","pt":"Qual é o limite da isenção de impostos?","ar":"ما هو حد الإعفاء الجمركي؟"},
   "a_ko":"해외에서 산 물품은 미화 팔백 달러까지 면세이며, 초과분은 세관에 신고하셔야 합니다.",
   "a":{"en":"Goods bought overseas are duty-free up to 800 US dollars, and anything over that must be declared to customs.","zh":"在境外购买的物品在800美元以内免税，超出部分必须向海关申报。","ja":"海外で購入した品物は800米ドルまで免税で、それを超える分は税関に申告する必要があります。","es":"Los productos comprados en el extranjero están libres de impuestos hasta 800 dólares estadounidenses, y lo que exceda debe declararse en la aduana.","fr":"Les articles achetés à l'étranger sont exonérés jusqu'à 800 dollars américains, et tout dépassement doit être déclaré à la douane.","pt":"Produtos comprados no exterior são isentos de impostos até 800 dólares americanos, e o que exceder deve ser declarado à alfândega.","ar":"البضائع المشتراة في الخارج معفاة من الرسوم حتى 800 دولار أمريكي، وما زاد عن ذلك يجب التصريح به للجمارك."}},
]

# ---- dataset.json ----
def write_dataset():
    scenarios = []
    for s in DATA:
        scenarios.append({
            "id": s["id"], "category": s["cat"], "difficulty": s["diff"], "traps": s["traps"],
            "question": {"ko_ref": s["q_ko"], "src": s["q"]},
            "answer": {"ko_src": s["a_ko"], "ref": s["a"]},
        })
    out = {
        "meta": {
            "version": 2,
            "langs": FOREIGN,
            "note": "30개 시나리오(난이도 하·중·상 각 10). 항공사명 포함 10문항, 버스터미널 1문항, 항공편 답변 일부는 '항공사 카운터 문의'로 유도. "
                    "question.src=외국인 발화 원문(STT 정답, 언어별), question.ko_ref=한국어 모범번역. answer.ko_src=한국인 발화 원문, answer.ref=언어별 모범번역. "
                    "⚠️ es/fr/pt/ar 은 초안 — 원어민 검수 후 확정할 것. (eval/gen_dataset.py 에서 생성)",
        },
        "trap_vocab": ["proper_noun", "gate_code", "flight_code", "number", "currency", "time", "procedure", "homophone"],
        "scenarios": scenarios,
    }
    p = os.path.join(HERE, "dataset.json")
    with open(p, "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)
    print("wrote", p, "-", len(scenarios), "scenarios")

# ---- 언어별 워드파일 ----
def write_docx():
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement

    def set_rtl(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        pPr.append(OxmlElement('w:bidi'))

    def add_line(doc, text, is_ar=False, bold=False, prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        if prefix:
            r = p.add_run(prefix); r.bold = True
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(11)
        if is_ar:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p)

    def add_item(doc, n, entries):  # entries: list of (lang, text, bold)
        for i, (lang, text, bold) in enumerate(entries):
            add_line(doc, text, is_ar=(lang == 'ar'), bold=bold, prefix=(f"{n}. " if i == 0 else None))
        g = doc.add_paragraph(); g.paragraph_format.space_after = Pt(2)

    outdir = os.path.join(HERE, "wordfiles")
    os.makedirs(outdir, exist_ok=True)
    for L in FOREIGN:
        doc = Document()
        doc.add_heading(f"KAC 통역 평가 · {LANG_DISP[L]} ({LANG_KO[L]}) 세트", level=0)
        note = doc.add_paragraph(
            "질문: 원어 → 영어 → 한국어 순 / 답변: 한국어 → 영어 → 원어 순. 굵은 줄이 낭독할 문장입니다. "
            "(원어=" + LANG_KO[L] + ")")
        note.runs[0].italic = True

        # 질문
        doc.add_heading("질문 (Questions) — 외국인 발화", level=1)
        for diff in DIFF_ORDER:
            doc.add_heading(f"난이도 {DIFF_KO[diff]}", level=2)
            n = 0
            for s in DATA:
                if s["diff"] != diff:
                    continue
                n += 1
                entries = [(L, s["q"][L], True)]      # 원어(낭독) — 굵게
                if L != 'en':
                    entries.append(('en', s["q"]['en'], False))  # 영어 참조
                entries.append(('ko', s["q_ko"], False))          # 한국어 참조
                add_item(doc, n, entries)

        # 답변
        doc.add_heading("답변 (Answers) — 한국인(안내원) 발화", level=1)
        for diff in DIFF_ORDER:
            doc.add_heading(f"난이도 {DIFF_KO[diff]}", level=2)
            n = 0
            for s in DATA:
                if s["diff"] != diff:
                    continue
                n += 1
                entries = [('ko', s["a_ko"], True)]   # 한국어(낭독) — 굵게
                if L != 'en':
                    entries.append(('en', s["a"]['en'], False))   # 영어 참조
                entries.append((L, s["a"][L], False))              # 원어 참조
                add_item(doc, n, entries)

        fn = os.path.join(outdir, f"KAC_평가_{LANG_KO[L]}.docx")
        doc.save(fn)
        print("wrote", fn)

if __name__ == "__main__":
    write_dataset()
    write_docx()
    # 무결성 점검
    assert len(DATA) == 30
    for diff in DIFF_ORDER:
        c = sum(1 for s in DATA if s["diff"] == diff)
        assert c == 10, f"{diff}={c}"
    air = sum(1 for s in DATA if any(k in s["cat"] for k in ["항공사"]) or "proper_noun" in s["traps"] and s["cat"] == "항공사")
    ids = [s["id"] for s in DATA]
    assert len(ids) == len(set(ids)), "중복 id"
    for s in DATA:
        for lg in FOREIGN:
            assert s["q"].get(lg) and s["a"].get(lg), f"{s['id']} 누락 {lg}"
    print("OK — 30 scenarios, 하/중/상 각 10, 7개 언어 완비")
